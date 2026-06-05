#!/usr/bin/env python3
from __future__ import annotations

import argparse
import csv
import hashlib
import json
import os
import re
import sqlite3
import subprocess
import sys
import zipfile
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

ROOT = Path.cwd()
MEMORY_ROOT = ROOT / ".yeyo-memory"
CONFIG_PATH = MEMORY_ROOT / "config.json"
VENDOR_PATH = MEMORY_ROOT / "vendor"

if VENDOR_PATH.exists():
    sys.path.insert(0, str(VENDOR_PATH))

try:
    from docx import Document
except Exception:
    Document = None

try:
    from openpyxl import load_workbook
except Exception:
    load_workbook = None

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None


@dataclass
class SourceFile:
    rel_path: str
    size_bytes: int
    size_human: str
    ext: str
    top_dir: str
    depth: int
    mtime_iso: str


def load_config() -> dict[str, Any]:
    return json.loads(CONFIG_PATH.read_text(encoding="utf-8"))


def doc_id(rel_path: str) -> str:
    return hashlib.sha1(rel_path.encode("utf-8")).hexdigest()[:16]


def safe_text(value: Any) -> str:
    if value is None:
        return ""
    text = str(value)
    return re.sub(r"\s+", " ", text).strip()


def normalize_text(text: str) -> str:
    text = text.replace("\x00", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def infer_title(rel_path: str) -> str:
    name = Path(rel_path).stem
    name = re.sub(r"[_]+", " ", name)
    name = re.sub(r"\s+", " ", name).strip()
    return name


def infer_doc_code(rel_path: str) -> str | None:
    name = Path(rel_path).name
    match = re.search(r"\b\d{4}[-_][A-Z]{3}[-_][A-Z]{3}[-_]\d{2,3}[-_]\d{2,3}[-_][A-Z0-9]{4}\b", name, re.I)
    if match:
        return match.group(0).replace("_", "-")
    match = re.search(r"\b\d{4}[-_][A-Z]{3}[-_][A-Z]{3}[-_][A-Z0-9-]{6,}\b", name, re.I)
    if match:
        return match.group(0).replace("_", "-")
    match = re.search(r"\b[A-Z]{2,4}[-_][A-Z]{2,4}[-_][A-Z]{2,4}[-_]\d{3}[-_]\d{2}\b", name, re.I)
    return match.group(0).replace("_", "-") if match else None


def infer_revision(rel_path: str) -> str | None:
    name = Path(rel_path).name
    patterns = [
        r"\bRev[ ._-]?([A-Z]?\d{1,3}|[A-Z])\b",
        r"\bR([A-Z]?\d{1,3})\b",
        r"\brev([A-Z]?\d{1,3})\b",
    ]
    for pattern in patterns:
        match = re.search(pattern, name, re.I)
        if match:
            return match.group(0)
    return None


def load_sources(index_path: Path) -> list[SourceFile]:
    with index_path.open(newline="", encoding="utf-8") as handle:
        reader = csv.DictReader(handle)
        return [
            SourceFile(
                rel_path=row["path"],
                size_bytes=int(row["size_bytes"]),
                size_human=row["size_human"],
                ext=row["ext"].lower(),
                top_dir=row["top_dir"],
                depth=int(row["depth"]),
                mtime_iso=row["mtime_iso"],
            )
            for row in reader
        ]


def extract_pdf(path: Path) -> tuple[str, dict[str, Any], str]:
    if PdfReader is None:
        return "", {}, "missing_dependency"
    try:
        reader = PdfReader(str(path))
        metadata = {
            "pages": len(reader.pages),
            "pdf_metadata": {str(k).lstrip("/"): safe_text(v) for k, v in (reader.metadata or {}).items()},
            "encrypted": bool(reader.is_encrypted),
        }
        if reader.is_encrypted:
            try:
                decrypt_result = reader.decrypt("")
                metadata["empty_password_decrypt_result"] = str(decrypt_result)
            except Exception as exc:
                metadata["decrypt_error"] = str(exc)
                return "", metadata, "encrypted"
            if not decrypt_result:
                return "", metadata, "encrypted"
        parts: list[str] = []
        for index, page in enumerate(reader.pages, start=1):
            try:
                page_text = page.extract_text() or ""
            except Exception as exc:
                page_text = f"\n[page {index} extraction_error: {exc}]\n"
            if page_text.strip():
                parts.append(f"\n\n[page {index}]\n{page_text}")
        return normalize_text("\n".join(parts)), metadata, "ok"
    except Exception as exc:
        return "", {"error": str(exc)}, "error"


def extract_docx(path: Path) -> tuple[str, dict[str, Any], str]:
    if Document is None:
        return "", {}, "missing_dependency"
    try:
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        table_rows: list[str] = []
        for table in doc.tables:
            for row in table.rows:
                cells = [safe_text(cell.text) for cell in row.cells]
                if any(cells):
                    table_rows.append(" | ".join(cells))
        text = "\n".join(paragraphs + table_rows)
        return normalize_text(text), {"paragraphs": len(paragraphs), "table_rows": len(table_rows)}, "ok"
    except Exception as exc:
        return "", {"error": str(exc)}, "error"


def extract_xlsx(path: Path) -> tuple[str, dict[str, Any], str]:
    if load_workbook is None:
        return "", {}, "missing_dependency"
    try:
        workbook = load_workbook(str(path), read_only=True, data_only=True)
        parts: list[str] = []
        sheet_stats: list[dict[str, Any]] = []
        for sheet in workbook.worksheets:
            rows_seen = 0
            non_empty = 0
            parts.append(f"\n\n[sheet {sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                rows_seen += 1
                values = [safe_text(value) for value in row]
                values = [value for value in values if value]
                if values:
                    non_empty += 1
                    parts.append(" | ".join(values))
            sheet_stats.append({"name": sheet.title, "rows_seen": rows_seen, "non_empty_rows": non_empty})
        workbook.close()
        return normalize_text("\n".join(parts)), {"sheets": sheet_stats}, "ok"
    except Exception as exc:
        return "", {"error": str(exc)}, "error"


def extract_doc_with_textutil(path: Path) -> tuple[str, dict[str, Any], str]:
    try:
        result = subprocess.run(
            ["textutil", "-convert", "txt", "-stdout", str(path)],
            check=False,
            capture_output=True,
            text=True,
            timeout=60,
        )
        metadata = {"returncode": result.returncode}
        if result.returncode != 0:
            metadata["stderr"] = result.stderr.strip()[:1000]
            return "", metadata, "error"
        return normalize_text(result.stdout), metadata, "ok"
    except Exception as exc:
        return "", {"error": str(exc)}, "error"


def inspect_archive(path: Path) -> tuple[str, dict[str, Any], str]:
    if path.suffix.lower() != ".zip":
        return "", {"note": "archive listing not supported for this extension"}, "metadata_only"
    try:
        with zipfile.ZipFile(path) as archive:
            infos = archive.infolist()
            names = [info.filename for info in infos[:500]]
            metadata = {
                "archive_entries": len(infos),
                "archive_sample": names,
                "archive_total_uncompressed_bytes": sum(info.file_size for info in infos),
            }
            return "\n".join(names), metadata, "archive_listed"
    except Exception as exc:
        return "", {"error": str(exc)}, "error"


def extract_text(source: SourceFile, config: dict[str, Any]) -> tuple[str, dict[str, Any], str]:
    path = ROOT / source.rel_path
    if source.ext == "pdf":
        return extract_pdf(path)
    if source.ext == "docx":
        return extract_docx(path)
    if source.ext in {"xlsx", "xlsm"}:
        return extract_xlsx(path)
    if source.ext == "doc":
        return extract_doc_with_textutil(path)
    if source.ext in set(config["archive_extensions"]):
        return inspect_archive(path)
    return "", {}, "metadata_only"


def chunk_text(text: str, chunk_chars: int, overlap: int) -> list[str]:
    if not text:
        return []
    chunks = []
    start = 0
    text_len = len(text)
    while start < text_len:
        end = min(start + chunk_chars, text_len)
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end == text_len:
            break
        start = max(0, end - overlap)
    return chunks


def write_json(path: Path, data: dict[str, Any]) -> None:
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def init_db(db_path: Path) -> sqlite3.Connection:
    db_path.parent.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(db_path)
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS documents (
            id TEXT PRIMARY KEY,
            path TEXT NOT NULL,
            top_dir TEXT NOT NULL,
            ext TEXT NOT NULL,
            size_bytes INTEGER NOT NULL,
            size_human TEXT NOT NULL,
            mtime_iso TEXT NOT NULL,
            title TEXT,
            doc_code TEXT,
            revision TEXT,
            status TEXT NOT NULL,
            text_chars INTEGER NOT NULL,
            chunks INTEGER NOT NULL,
            token_estimate INTEGER NOT NULL,
            card_path TEXT NOT NULL,
            extracted_path TEXT
        )
        """
    )
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS chunks (
            id TEXT PRIMARY KEY,
            document_id TEXT NOT NULL,
            chunk_index INTEGER NOT NULL,
            path TEXT NOT NULL,
            text TEXT NOT NULL,
            FOREIGN KEY(document_id) REFERENCES documents(id)
        )
        """
    )
    conn.execute("CREATE VIRTUAL TABLE IF NOT EXISTS chunks_fts USING fts5(id UNINDEXED, path UNINDEXED, text)")
    conn.commit()
    return conn


def upsert_document(conn: sqlite3.Connection, card: dict[str, Any]) -> None:
    conn.execute(
        """
        INSERT OR REPLACE INTO documents (
            id, path, top_dir, ext, size_bytes, size_human, mtime_iso, title,
            doc_code, revision, status, text_chars, chunks, token_estimate,
            card_path, extracted_path
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            card["id"],
            card["path"],
            card["top_dir"],
            card["ext"],
            card["size_bytes"],
            card["size_human"],
            card["mtime_iso"],
            card["title"],
            card.get("doc_code"),
            card.get("revision"),
            card["status"],
            card["text_chars"],
            card["chunks"],
            card["token_estimate"],
            card["card_path"],
            card.get("extracted_path"),
        ),
    )


def replace_chunks(conn: sqlite3.Connection, document_id: str, rows: list[dict[str, Any]]) -> None:
    conn.execute("DELETE FROM chunks WHERE document_id = ?", (document_id,))
    conn.execute("DELETE FROM chunks_fts WHERE id IN (SELECT id FROM chunks_fts WHERE id LIKE ?)", (f"{document_id}:%",))
    for row in rows:
        conn.execute(
            "INSERT OR REPLACE INTO chunks (id, document_id, chunk_index, path, text) VALUES (?, ?, ?, ?, ?)",
            (row["id"], row["document_id"], row["chunk_index"], row["path"], row["text"]),
        )
        conn.execute("INSERT INTO chunks_fts (id, path, text) VALUES (?, ?, ?)", (row["id"], row["path"], row["text"]))


def export_chunks_jsonl(conn: sqlite3.Connection, path: Path) -> None:
    with path.open("w", encoding="utf-8") as handle:
        rows = conn.execute(
            "SELECT id, document_id, chunk_index, path, text FROM chunks ORDER BY path, chunk_index"
        )
        for row in rows:
            handle.write(
                json.dumps(
                    {
                        "id": row[0],
                        "document_id": row[1],
                        "chunk_index": row[2],
                        "path": row[3],
                        "text": row[4],
                    },
                    ensure_ascii=False,
                )
                + "\n"
            )


def process_source(source: SourceFile, config: dict[str, Any]) -> tuple[dict[str, Any], list[dict[str, Any]]]:
    did = doc_id(source.rel_path)
    text, metadata, extraction_status = extract_text(source, config)
    min_text = int(config["min_text_chars"])
    status = extraction_status
    if extraction_status == "ok" and len(text) < min_text:
        status = "needs_ocr_or_empty"
    if extraction_status == "metadata_only":
        status = "metadata_only"
    chunks = chunk_text(text, int(config["chunk_chars"]), int(config["chunk_overlap"])) if len(text) >= min_text else []

    extracted_rel = None
    if text:
        extracted_path = MEMORY_ROOT / "extracted" / f"{did}.txt"
        extracted_path.write_text(text + "\n", encoding="utf-8")
        extracted_rel = str(extracted_path.relative_to(ROOT))

    chunk_rows = [
        {
            "id": f"{did}:{index:04d}",
            "document_id": did,
            "chunk_index": index,
            "path": source.rel_path,
            "text": chunk,
        }
        for index, chunk in enumerate(chunks)
    ]

    card_rel = str((MEMORY_ROOT / "cards" / f"{did}.json").relative_to(ROOT))
    card = {
        "id": did,
        "path": source.rel_path,
        "top_dir": source.top_dir,
        "ext": source.ext,
        "size_bytes": source.size_bytes,
        "size_human": source.size_human,
        "mtime_iso": source.mtime_iso,
        "title": infer_title(source.rel_path),
        "doc_code": infer_doc_code(source.rel_path),
        "revision": infer_revision(source.rel_path),
        "status": status,
        "text_chars": len(text),
        "token_estimate": max(1, len(text) // 4) if text else 0,
        "chunks": len(chunks),
        "card_path": card_rel,
        "extracted_path": extracted_rel,
        "metadata": metadata,
        "processed_at": datetime.now(timezone.utc).isoformat(),
    }
    return card, chunk_rows


def read_existing_state(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {}
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {}


def write_reports(cards: list[dict[str, Any]], processed: int, skipped: int) -> None:
    reports = MEMORY_ROOT / "reports"
    reports.mkdir(parents=True, exist_ok=True)
    by_status: dict[str, int] = {}
    by_ext: dict[str, int] = {}
    by_top: dict[str, dict[str, int]] = {}
    for card in cards:
        by_status[card["status"]] = by_status.get(card["status"], 0) + 1
        by_ext[card["ext"]] = by_ext.get(card["ext"], 0) + 1
        bucket = by_top.setdefault(card["top_dir"], {"files": 0, "text_chars": 0, "chunks": 0})
        bucket["files"] += 1
        bucket["text_chars"] += card["text_chars"]
        bucket["chunks"] += card["chunks"]

    summary = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "cards": len(cards),
        "processed_this_run": processed,
        "skipped_unchanged": skipped,
        "by_status": dict(sorted(by_status.items())),
        "by_ext": dict(sorted(by_ext.items())),
        "by_top_dir": dict(sorted(by_top.items())),
    }
    write_json(reports / "processing-summary.json", summary)

    lines = [
        "# Processing summary",
        "",
        f"- Generated: {summary['generated_at']}",
        f"- Cards: {summary['cards']}",
        f"- Processed this run: {processed}",
        f"- Skipped unchanged: {skipped}",
        "",
        "## Status",
        "",
        "| Status | Files |",
        "|---|---:|",
        *[f"| {status} | {count} |" for status, count in sorted(by_status.items())],
        "",
        "## Top folders",
        "",
        "| Folder | Files | Text chars | Chunks |",
        "|---|---:|---:|---:|",
        *[
            f"| {top} | {data['files']} | {data['text_chars']} | {data['chunks']} |"
            for top, data in sorted(by_top.items())
        ],
    ]
    (reports / "processing-summary.md").write_text("\n".join(lines) + "\n", encoding="utf-8")

    report_columns = [
        "status",
        "path",
        "ext",
        "top_dir",
        "size_human",
        "text_chars",
        "chunks",
        "doc_code",
        "revision",
        "card_path",
        "extracted_path",
    ]
    with (reports / "documents-by-status.csv").open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=report_columns)
        writer.writeheader()
        for card in sorted(cards, key=lambda item: (item["status"], item["top_dir"], item["path"])):
            writer.writerow({column: card.get(column, "") for column in report_columns})

    for status, filename in [("needs_ocr_or_empty", "needs-ocr.csv"), ("error", "errors.csv"), ("metadata_only", "metadata-only.csv")]:
        with (reports / filename).open("w", newline="", encoding="utf-8") as handle:
            writer = csv.DictWriter(handle, fieldnames=report_columns)
            writer.writeheader()
            for card in sorted((card for card in cards if card["status"] == status), key=lambda item: item["path"]):
                writer.writerow({column: card.get(column, "") for column in report_columns})


def main() -> int:
    parser = argparse.ArgumentParser(description="Build local document memory for the Yeyo folder.")
    parser.add_argument("--top-dir", help="Only process documents under a top-level folder.")
    parser.add_argument("--limit", type=int, help="Only process the first N matching files.")
    parser.add_argument("--force", action="store_true", help="Reprocess unchanged files.")
    args = parser.parse_args()

    config = load_config()
    for key in ["cards", "extracted", "chunks", "reports"]:
        (ROOT / config["paths"][key]).mkdir(parents=True, exist_ok=True)

    sources = load_sources(ROOT / config["paths"]["source_index"])
    if args.top_dir:
        sources = [source for source in sources if source.top_dir == args.top_dir]
    if args.limit:
        sources = sources[: args.limit]

    state_path = MEMORY_ROOT / "state.json"
    state = read_existing_state(state_path)
    seen = state.get("files", {})
    conn = init_db(ROOT / config["paths"]["sqlite"])

    all_cards: list[dict[str, Any]] = []
    processed = 0
    skipped = 0
    for source in sources:
        did = doc_id(source.rel_path)
        card_path = MEMORY_ROOT / "cards" / f"{did}.json"
        unchanged = (
            not args.force
            and seen.get(source.rel_path, {}).get("size_bytes") == source.size_bytes
            and seen.get(source.rel_path, {}).get("mtime_iso") == source.mtime_iso
            and card_path.exists()
        )
        if unchanged:
            card = json.loads(card_path.read_text(encoding="utf-8"))
            skipped += 1
        else:
            card, chunks = process_source(source, config)
            write_json(card_path, card)
            replace_chunks(conn, did, chunks)
            seen[source.rel_path] = {
                "id": did,
                "size_bytes": source.size_bytes,
                "mtime_iso": source.mtime_iso,
                "status": card["status"],
            }
            processed += 1

        upsert_document(conn, card)
        all_cards.append(card)

        if processed and processed % 100 == 0:
            conn.commit()
            print(f"processed={processed} skipped={skipped}", file=sys.stderr)

    state["updated_at"] = datetime.now(timezone.utc).isoformat()
    state["files"] = seen
    write_json(state_path, state)
    conn.commit()
    export_chunks_jsonl(conn, MEMORY_ROOT / "chunks" / "chunks.jsonl")
    conn.close()
    write_reports(all_cards, processed, skipped)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
